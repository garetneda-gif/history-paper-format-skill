#!/usr/bin/env python3
"""
format_docx.py — 《历史研究》期刊排版主入口

用法:
    python3 format_docx.py --input 论文.docx --output 排版后.docx
    python3 format_docx.py --input 论文.docx --output 排版后.docx --no-backup

退出码:
    0  成功
    1  参数错误 / 文件不存在 / 格式错误
"""

import argparse
import sys
from pathlib import Path

# ── 将 lib/ 加入查找路径 ────────────────────────────────────────────────────
sys.path.insert(0, str(Path(__file__).parent))

from docx import Document

from lib.specs import DEFAULT_SPEC, STYLE_NAMES
from lib.io_utils import resolve_input, resolve_output, backup_input
from lib.style_factory import ensure_paragraph_styles, apply_base_page_setup
from lib.font_utils import set_run_fonts, set_style_fonts
from lib.paragraph_rules import classify_paragraph, apply_paragraph_style
from lib.footnote_ooxml import set_footnote_restart_each_page


# ── 样式名 → 东亚字体映射 ────────────────────────────────────────────────────
_STYLE_EAST_FONT: dict[str, str] = {
    STYLE_NAMES["body"]:           DEFAULT_SPEC.font_body_east,
    STYLE_NAMES["quote"]:          DEFAULT_SPEC.font_quote_east,
    STYLE_NAMES["footnote"]:       DEFAULT_SPEC.font_footnote_east,
    STYLE_NAMES["title_main"]:     DEFAULT_SPEC.font_title_east,
    STYLE_NAMES["subtitle"]:       DEFAULT_SPEC.font_subtitle_east,
    STYLE_NAMES["abstract_label"]: DEFAULT_SPEC.font_abstract_label_east,
    STYLE_NAMES["abstract_text"]:  DEFAULT_SPEC.font_abstract_text_east,
    STYLE_NAMES["section_l2"]:     DEFAULT_SPEC.font_section_l2_east,
}


def _apply_style_level_fonts(doc) -> None:
    """在命名样式层面设置拉丁 + 东亚字体，为所有段落提供继承基础。"""
    latin = DEFAULT_SPEC.font_latin
    for style_name, east in _STYLE_EAST_FONT.items():
        if style_name in [s.name for s in doc.styles]:
            set_style_fonts(doc.styles[style_name], latin, east)


def _apply_run_level_fonts(doc) -> None:
    """对每个 run 直写字体，确保中西文分离不被旧 run 属性覆盖。"""
    latin = DEFAULT_SPEC.font_latin
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else STYLE_NAMES["body"]
        east = _STYLE_EAST_FONT.get(style_name, DEFAULT_SPEC.font_body_east)
        for run in para.runs:
            set_run_fonts(run, latin, east)


def _process_paragraphs(doc) -> None:
    """分类并应用段落样式。"""
    for para in doc.paragraphs:
        style_name = classify_paragraph(para)
        apply_paragraph_style(para, style_name)


def format_document(input_path: Path, output_path: Path) -> None:
    """执行完整排版流程。"""
    doc = Document(str(input_path))

    # 1. 页面设置（边距）
    apply_base_page_setup(doc)

    # 2. 注册 / 确保 8 个命名样式存在
    ensure_paragraph_styles(doc)

    # 3. 样式层字体（继承基础）
    _apply_style_level_fonts(doc)

    # 4. 段落分类 + 样式应用
    _process_paragraphs(doc)

    # 5. run 层字体（直写，覆盖旧属性）
    _apply_run_level_fonts(doc)

    # 6. 保存（必须先 save，OOXML 补丁在 zip 层操作）
    doc.save(str(output_path))

    # 7. 脚注每页重排（OOXML 直改 settings.xml）
    set_footnote_restart_each_page(str(output_path))


def _build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        description="将 Word 文档按《历史研究》期刊规范排版",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "示例:\n"
            "  python3 format_docx.py --input 论文.docx --output 排版后.docx\n"
            "  python3 format_docx.py -i 论文.docx -o 排版后.docx --no-backup\n"
        ),
    )
    p.add_argument("-i", "--input",  required=True, help="输入 .docx 文件路径")
    p.add_argument("-o", "--output", required=True, help="输出 .docx 文件路径")
    p.add_argument(
        "--no-backup",
        action="store_true",
        default=False,
        help="不备份原始文件（默认会生成 .bak.docx）",
    )
    return p


def main() -> None:
    parser = _build_parser()
    args = parser.parse_args()

    input_path  = resolve_input(args.input)
    output_path = resolve_output(args.output, input_path)

    if not args.no_backup:
        bak = backup_input(input_path)
        print(f"[备份] {bak}")

    print(f"[开始] {input_path} → {output_path}")
    format_document(input_path, output_path)
    print(f"[完成] 排版输出: {output_path}")


if __name__ == "__main__":
    main()
